# import tkinter as tk
# from tkinter import ttk, filedialog, scrolledtext, messagebox
# import os
# import docx
# from openpyxl import load_workbook
# import pandas as pd
# import requests
# import json
# import re
# from dotenv import load_dotenv
# load_dotenv()

# QWEN_API_KEY = os.getenv("QWEN_API_KEY")

# class FileEditorAgent:
#     def __init__(self, root):
#         self.root = root
#         self.root.title("AI File Editor Agent")
#         self.root.geometry("1000x700")
        
#         # Variables
#         self.current_folder = ""
#         self.current_file = ""
#         self.file_content = ""
#         self.api_endpoint = "https://openrouter.ai/api/v1/chat/completions"
#         self.model = "deepseek/deepseek-r1-0528-qwen3-8b:free"
#         self.ai_suggestion = None  # Store the latest AI suggestion
        
#         # Create UI
#         self.create_widgets()
    
#     def create_widgets(self):
#         # Main panes
#         main_pane = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
#         main_pane.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
#         # Left pane (file browser and editor)
#         left_frame = ttk.Frame(main_pane)
#         main_pane.add(left_frame, weight=2)
        
#         # Right pane (AI assistant)
#         right_frame = ttk.Frame(main_pane)
#         main_pane.add(right_frame, weight=1)
        
#         # Folder selection
#         folder_frame = ttk.Frame(left_frame, padding=5)
#         folder_frame.pack(fill=tk.X)
        
#         ttk.Label(folder_frame, text="Folder:").pack(side=tk.LEFT)
#         self.folder_label = ttk.Label(folder_frame, text="No folder selected", width=40)
#         self.folder_label.pack(side=tk.LEFT, padx=5)
        
#         ttk.Button(
#             folder_frame, 
#             text="Browse", 
#             command=self.select_folder
#         ).pack(side=tk.RIGHT)
        
#         # File list
#         file_frame = ttk.LabelFrame(left_frame, text="Files", padding=5)
#         file_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
#         self.file_tree = ttk.Treeview(file_frame, columns=("size",), show="tree", height=10)
#         self.file_tree.heading("#0", text="Name")
#         self.file_tree.heading("size", text="Size")
#         self.file_tree.pack(fill=tk.BOTH, expand=True, side=tk.LEFT)
        
#         scrollbar = ttk.Scrollbar(file_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
#         scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
#         self.file_tree.configure(yscrollcommand=scrollbar.set)
#         self.file_tree.bind("<<TreeviewSelect>>", self.on_file_select)
        
#         # Editor
#         editor_frame = ttk.LabelFrame(left_frame, text="Editor", padding=5)
#         editor_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
#         self.editor = scrolledtext.ScrolledText(editor_frame, wrap=tk.WORD, height=15)
#         self.editor.pack(fill=tk.BOTH, expand=True)
#         self.editor.config(state=tk.DISABLED)
        
#         # Buttons
#         button_frame = ttk.Frame(left_frame, padding=5)
#         button_frame.pack(fill=tk.X)
        
#         self.save_btn = ttk.Button(
#             button_frame,
#             text="Save Changes",
#             command=self.save_file,
#             state=tk.DISABLED
#         )
#         self.save_btn.pack(side=tk.RIGHT, padx=5)
        
#         # AI Assistant Panel
#         ai_frame = ttk.LabelFrame(right_frame, text="AI Assistant", padding=10)
#         ai_frame.pack(fill=tk.BOTH, expand=True)
        
#         ttk.Label(ai_frame, text="Ask about your file:").pack(anchor=tk.W)
        
#         self.user_input = ttk.Entry(ai_frame)
#         self.user_input.pack(fill=tk.X, pady=5)
#         self.user_input.bind("<Return>", self.ask_ai)
        
#         ttk.Button(
#             ai_frame, 
#             text="Ask AI", 
#             command=self.ask_ai
#         ).pack(anchor=tk.E, pady=5)
        
#         self.ai_response = scrolledtext.ScrolledText(
#             ai_frame, 
#             wrap=tk.WORD, 
#             height=20,
#             state=tk.DISABLED
#         )
#         self.ai_response.pack(fill=tk.BOTH, expand=True)
        
#         # Apply AI Suggestion Button
#         self.apply_btn = ttk.Button(
#             ai_frame,
#             text="Apply Suggestion",
#             command=self.apply_ai_suggestion,
#             state=tk.DISABLED
#         )
#         self.apply_btn.pack(anchor=tk.E, pady=5)
    
#     def select_folder(self):
#         folder_selected = filedialog.askdirectory()
#         if folder_selected:
#             self.current_folder = folder_selected
#             self.folder_label.config(text=folder_selected)
#             self.load_file_list()
    
#     def load_file_list(self):
#         self.file_tree.delete(*self.file_tree.get_children())
#         for item in os.listdir(self.current_folder):
#             full_path = os.path.join(self.current_folder, item)
#             if os.path.isfile(full_path):
#                 size = os.path.getsize(full_path)
#                 self.file_tree.insert("", "end", text=item, values=(f"{size} bytes",))
    
#     def on_file_select(self, event):
#         selected = self.file_tree.focus()
#         if not selected:
#             return
        
#         filename = self.file_tree.item(selected, "text")
#         self.current_file = os.path.join(self.current_folder, filename)
        
#         try:
#             if filename.endswith(('.txt', '.py', '.js', '.html', '.css', '.json', '.xml')):
#                 self.load_text_file()
#             elif filename.endswith('.docx'):
#                 self.load_docx_file()
#             elif filename.endswith(('.xlsx', '.xls')):
#                 self.load_excel_file()
#             else:
#                 messagebox.showwarning("Unsupported Format", "File type not supported")
#                 return
            
#             self.editor.config(state=tk.NORMAL)
#             self.save_btn.config(state=tk.NORMAL)
#             self.ai_suggestion = None  # Reset previous suggestion
#             self.apply_btn.config(state=tk.DISABLED)
#         except Exception as e:
#             messagebox.showerror("Error", f"Failed to load file: {str(e)}")
    
#     def load_text_file(self):
#         with open(self.current_file, 'r', encoding='utf-8') as file:
#             self.file_content = file.read()
#         self.editor.delete(1.0, tk.END)
#         self.editor.insert(tk.END, self.file_content)
    
#     def load_docx_file(self):
#         doc = docx.Document(self.current_file)
#         self.file_content = "\n".join([para.text for para in doc.paragraphs])
#         self.editor.delete(1.0, tk.END)
#         self.editor.insert(tk.END, self.file_content)
    
#     def load_excel_file(self):
#         wb = load_workbook(self.current_file)
#         sheet = wb.active
#         df = pd.DataFrame(sheet.values)
#         self.file_content = df.to_csv(index=False, header=False)
#         self.editor.delete(1.0, tk.END)
#         self.editor.insert(tk.END, self.file_content)
    
#     def save_file(self):
#         new_content = self.editor.get(1.0, tk.END)[:-1]  # Remove trailing newline
        
#         try:
#             if self.current_file.endswith(('.txt', '.py', '.js', '.html', '.css', '.json', '.xml')):
#                 with open(self.current_file, 'w', encoding='utf-8') as file:
#                     file.write(new_content)
            
#             elif self.current_file.endswith('.docx'):
#                 doc = docx.Document()
#                 for line in new_content.split('\n'):
#                     doc.add_paragraph(line)
#                 doc.save(self.current_file)
            
#             elif self.current_file.endswith(('.xlsx', '.xls')):
#                 from io import StringIO
#                 df = pd.read_csv(StringIO(new_content))
#                 df.to_excel(self.current_file, index=False, header=False)
            
#             messagebox.showinfo("Success", "File saved successfully!")
#         except Exception as e:
#             messagebox.showerror("Error", f"Failed to save file: {str(e)}")
    
#     def ask_ai(self, event=None):
#         if not self.current_file:
#             messagebox.showwarning("No File", "Please select a file first")
#             return
            
#         user_query = self.user_input.get()
#         if not user_query:
#             return
            
#         # Get current file content
#         current_content = self.editor.get(1.0, tk.END)
#         filename = os.path.basename(self.current_file)
        
#         # Prepare prompt for the AI
#         prompt = (
#             f"You are an expert file assistant. The user is working with a file named '{filename}'. "
#             f"Here is the current content of the file:\n\n"
#             f"```\n{current_content}\n```\n\n"
#             f"User request: {user_query}\n\n"
#             "If the request requires changes to the file, provide ONLY the modified content and old_content which was replaced by modified_content in a code block. "
#             "Otherwise, provide a helpful response explaining how to accomplish the request."
#         )
        
#         # Prepare schema for structured response
#         schema = {
#             "type": "object",
#             "properties": {
#                 "explanation": {"type": "string"},
#                 "modified_content": {"type": "string"},
#                 "old_content": {"type": "string"},
#             }
#         }
        
#         try:
#             # Call API
#             res = requests.post(
#                 self.api_endpoint,
#                 headers={
#                     "Authorization": f"Bearer {QWEN_API_KEY}",
#                     "Content-Type": "application/json"
#                 },
#                 json={
#                     "model": self.model,
#                     "messages": [
#                         {"role": "system", "content": "You are an expert file assistant."},
#                         {"role": "user", "content": prompt}
#                     ],
#                     "response_format": {
#                         "type": "json_object",
#                         "schema": schema
#                     },
#                     "temperature": 0.2
#                 },
#                 timeout=120
#             )
            
#             if res.status_code != 200:
#                 messagebox.showerror("API Error", f"API request failed: {res.status_code}\n{res.text}")
#                 return
                
#             response_data = res.json()
#             print("API Response:", json.dumps(response_data, indent=2))
            
#             # Extract the content from the response
#             ai_content = response_data["choices"][0]["message"]["content"]
            
#             try:
#                 # Parse the JSON content
#                 parsed_content = json.loads(ai_content)
#                 explanation = parsed_content.get("explanation", "No explanation provided.")
#                 modified_content = parsed_content.get("modified_content", "")
#                 old_content = parsed_content.get("old_content", "")
                
#                 # Store the modified content for later application
#                 self.ai_suggestion = modified_content
                
#                 # Format display text
#                 display_text = f"Explanation: {explanation}\n\n"
#                 if old_content:
#                     display_text += f"Old Content:\n{old_content}\n\n"
#                 if modified_content:
#                     display_text += f"Modified Content:\n{modified_content}"
                
#                 # Enable apply button if we have modified content
#                 if modified_content:
#                     self.apply_btn.config(state=tk.NORMAL)
#                 else:
#                     self.apply_btn.config(state=tk.DISABLED)
                    
#             except json.JSONDecodeError:
#                 # If response isn't JSON, use it as is
#                 display_text = ai_content
#                 self.ai_suggestion = None
#                 self.apply_btn.config(state=tk.DISABLED)
            
#             # Display AI response
#             self.ai_response.config(state=tk.NORMAL)
#             self.ai_response.delete(1.0, tk.END)
#             self.ai_response.insert(tk.END, display_text)
#             self.ai_response.config(state=tk.DISABLED)
                
#         except Exception as e:
#             messagebox.showerror("Error", f"Failed to get AI response: {str(e)}")
#             import traceback
#             traceback.print_exc()
    
#     def apply_ai_suggestion(self):
#         if self.ai_suggestion:
#             # Update editor with the AI suggestion
#             self.editor.delete(1.0, tk.END)
#             self.editor.insert(tk.END, self.ai_suggestion)
#             messagebox.showinfo("Suggestion Applied", "AI suggestion applied to editor. Click 'Save Changes' to save to file.")
#         else:
#             # Fallback to text extraction method
#             ai_text = self.ai_response.get(1.0, tk.END)
#             matches = re.findall(r'```(?:[^\n]*\n)?(.*?)```', ai_text, re.DOTALL)
            
#             if matches:
#                 new_content = matches[-1].strip()
#                 self.editor.delete(1.0, tk.END)
#                 self.editor.insert(tk.END, new_content)
#                 messagebox.showinfo("Suggestion Applied", "AI suggestion applied to editor.")
#             else:
#                 messagebox.showinfo("No Suggestion", "No suggestion found in AI response")

# if __name__ == "__main__":
#     root = tk.Tk()
#     app = FileEditorAgent(root)
#     root.mainloop()

import os
import docx
from openpyxl import load_workbook
import pandas as pd
import requests
import json
import re
from dotenv import load_dotenv

load_dotenv()

QWEN_API_KEY = os.getenv("QWEN_API_KEY")

class FileEditorAgent:
    def __init__(self):
        self.current_folder = ""
        self.current_file = ""
        self.file_content = ""
        self.api_endpoint = "https://openrouter.ai/api/v1/chat/completions"
        self.model = "deepseek/deepseek-r1-0528-qwen3-8b:free"
        self.ai_suggestion = None

    def load_file_list(self):
        if not self.current_folder:
            return []
        
        file_list = []
        for item in os.listdir(self.current_folder):
            full_path = os.path.join(self.current_folder, item)
            if os.path.isfile(full_path):
                size = os.path.getsize(full_path)
                file_list.append((item, f"{size} bytes"))
        return file_list

    def load_file_content(self, filename):
        self.current_file = os.path.join(self.current_folder, filename)
        
        if filename.endswith(('.txt', '.py', '.js', '.html', '.css', '.json', '.xml')):
            return self.load_text_file()
        elif filename.endswith('.docx'):
            return self.load_docx_file()
        elif filename.endswith(('.xlsx', '.xls')):
            return self.load_excel_file()
        else:
            raise ValueError("Unsupported file format")

    def load_text_file(self):
        with open(self.current_file, 'r', encoding='utf-8') as file:
            self.file_content = file.read()
        return self.file_content

    def load_docx_file(self):
        doc = docx.Document(self.current_file)
        self.file_content = "\n".join([para.text for para in doc.paragraphs])
        return self.file_content

    def load_excel_file(self):
        wb = load_workbook(self.current_file)
        sheet = wb.active
        df = pd.DataFrame(sheet.values)
        self.file_content = df.to_csv(index=False, header=False)
        return self.file_content

    def save_file(self, new_content):
        # Remove trailing newline added by Text widget
        new_content = new_content.rstrip('\n')
        
        if self.current_file.endswith(('.txt', '.py', '.js', '.html', '.css', '.json', '.xml')):
            with open(self.current_file, 'w', encoding='utf-8') as file:
                file.write(new_content)
        elif self.current_file.endswith('.docx'):
            doc = docx.Document()
            for line in new_content.split('\n'):
                doc.add_paragraph(line)
            doc.save(self.current_file)
        elif self.current_file.endswith(('.xlsx', '.xls')):
            from io import StringIO
            df = pd.read_csv(StringIO(new_content))
            df.to_excel(self.current_file, index=False, header=False)
        return True

    def ask_ai(self, file_content, filename, user_query):
        # Prepare prompt for the AI
        prompt = (
            f"You are an expert file assistant. The user is working with a file named '{filename}'. "
            f"Here is the current content of the file:\n\n"
            f"```\n{file_content}\n```\n\n"
            f"User request: {user_query}\n\n"
            "If the request requires changes to the file, provide ONLY the modified content and old_content which was replaced by modified_content in a code block. "
            "Otherwise, provide a helpful response explaining how to accomplish the request."
        )
        
        # Prepare schema for structured response
        schema = {
            "type": "object",
            "properties": {
                "explanation": {"type": "string"},
                "modified_content": {"type": "string"},
                "old_content": {"type": "string"},
            }
        }
        
        try:
            # Call API
            res = requests.post(
                self.api_endpoint,
                headers={
                    "Authorization": f"Bearer {QWEN_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": self.model,
                    "messages": [
                        {"role": "system", "content": "You are an expert file assistant."},
                        {"role": "user", "content": prompt}
                    ],
                    "response_format": {
                        "type": "json_object",
                        "schema": schema
                    },
                    "temperature": 0.2
                },
                timeout=120
            )
            
            if res.status_code != 200:
                return None, f"API request failed: {res.status_code}\n{res.text}"
                
            response_data = res.json()
            
            # Extract the content from the response
            ai_content = response_data["choices"][0]["message"]["content"]
            
            try:
                # Parse the JSON content
                parsed_content = json.loads(ai_content)
                explanation = parsed_content.get("explanation", "No explanation provided.")
                modified_content = parsed_content.get("modified_content", "")
                old_content = parsed_content.get("old_content", "")
                
                # Store the modified content for later application
                self.ai_suggestion = modified_content
                
                # Format display text
                display_text = f"Explanation: {explanation}\n\n"
                if old_content:
                    display_text += f"Old Content:\n{old_content}\n\n"
                if modified_content:
                    display_text += f"Modified Content:\n{modified_content}"
                
                return display_text, modified_content
                    
            except json.JSONDecodeError:
                # If response isn't JSON, use it as is
                return ai_content, None
            
        except Exception as e:
            return None, f"Failed to get AI response: {str(e)}"

    def apply_ai_suggestion(self):
        return self.ai_suggestion